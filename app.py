import streamlit as st
from docx import Document
import pdfplumber
import json
import os
import re
from openai import OpenAI

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

st.title("CV Auto Formatter")

cv_file = st.file_uploader("Upload CV (PDF ou DOCX)", type=["pdf", "docx"])
template_file = st.file_uploader("Upload Template DOCX", type=["docx"])


def extract_text(file):
    if file.name.lower().endswith(".pdf"):
        text = ""

        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

        return text

    elif file.name.lower().endswith(".docx"):
        doc = Document(file)
        return "\n".join([p.text for p in doc.paragraphs])

    return ""


def replace_text_in_paragraphs(doc, replacements):
    for p in doc.paragraphs:
        for key, value in replacements.items():
            if key in p.text:
                p.text = p.text.replace(key, str(value))


if st.button("Gerar CV"):
    if not cv_file or not template_file:
        st.error("Faz upload dos dois ficheiros")
    else:
        with st.spinner("A processar..."):
            cv_text = extract_text(cv_file)

            prompt = f"""
Extrai os seguintes campos do CV:

- initials
- experience
- availability
- english_level
- country
- skills
- summary
- work_experience

O campo work_experience deve ser uma lista com:
- company
- position
- start
- end
- description

Responde apenas em JSON válido, sem markdown.

CV:
{cv_text}
"""

            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "user", "content": prompt}
                ],
                temperature=0
            )

            raw = response.choices[0].message.content

            match = re.search(r"\{.*\}", raw, re.DOTALL)

            if not match:
                st.error("Erro ao interpretar resposta da IA")
                st.write(raw)
                st.stop()

            try:
                data = json.loads(match.group(0))
            except json.JSONDecodeError:
                st.error("A IA respondeu com JSON inválido")
                st.write(raw)
                st.stop()

            doc = Document(template_file)

            skills = data.get("skills", "")
            if isinstance(skills, list):
                skills = "; ".join(skills)

            replacements = {
                "[Candidates First and Last Initials]": data.get("initials", ""),
                "[Experience]": data.get("experience", ""),
                "[Availability]": data.get("availability", ""),
                "[English]": data.get("english_level", ""),
                "[Country]": data.get("country", ""),
                "[Skills]": skills,
                "[Summary]": data.get("summary", "")
            }

            replace_text_in_paragraphs(doc, replacements)

            work_experience = data.get("work_experience", [])

            if work_experience:
                exp = work_experience[0]

                exp_replacements = {
                    "[Company Name]": f"{exp.get('company', '')} - {exp.get('position', '')}",
                    "[start date]": f"{exp.get('start', '')} to {exp.get('end', '')}",
                    "[Description]": exp.get("description", "")
                }

                replace_text_in_paragraphs(doc, exp_replacements)

            output = "output.docx"
            doc.save(output)

            with open(output, "rb") as f:
                st.download_button(
                    "Download CV",
                    f,
                    file_name="CV_final.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
